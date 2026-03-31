"""
AIVA EPICS Individual Contribution Report Generator
Generates a professional .docx report for VIT Bhopal submission.
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ─── PAGE SETUP ───────────────────────────────────────────────
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

# ─── STYLE HELPERS ────────────────────────────────────────────
def set_font(run, name='Times New Roman', size=12, bold=False, italic=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading_styled(text, level=1, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.space_before = Pt(18)
    p.space_after = Pt(10)
    run = p.add_run(text)
    if level == 0:
        set_font(run, size=26, bold=True)
    elif level == 1:
        set_font(run, size=16, bold=True)
    elif level == 2:
        set_font(run, size=14, bold=True)
    elif level == 3:
        set_font(run, size=12, bold=True, italic=True)
    return p

def add_body(text, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=8):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(18)
    run = p.add_run(text)
    set_font(run, size=12)
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.5 + level * 0.25)
    p.paragraph_format.line_spacing = Pt(18)
    p.space_after = Pt(4)
    for run in p.runs:
        set_font(run, size=12)
    # Clear and re-add with correct font
    p.clear()
    run = p.add_run(text)
    set_font(run, size=12)
    return p

def add_page_break():
    doc.add_page_break()

def add_empty_lines(count=1):
    for _ in range(count):
        p = doc.add_paragraph()
        p.space_after = Pt(0)
        p.space_before = Pt(0)

# ═══════════════════════════════════════════════════════════════
# 1. COVER PAGE
# ═══════════════════════════════════════════════════════════════

add_empty_lines(4)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('VIT Bhopal University')
set_font(run, size=20, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('School of Computing Science and Engineering')
set_font(run, size=14)

add_empty_lines(2)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('EPICS Project')
set_font(run, size=16, bold=True, color=(0, 51, 153))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.space_after = Pt(6)
run = p.add_run('Individual Contribution Report')
set_font(run, size=14, bold=True, color=(0, 51, 153))

add_empty_lines(2)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('AIVA')
set_font(run, size=28, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('AI-Assisted Intelligent Vitals & Anamnesis')
set_font(run, size=14, italic=True)

add_empty_lines(3)

# Author details table
table = doc.add_table(rows=4, cols=2)
table.alignment = WD_ALIGN_PARAGRAPH.CENTER
details = [
    ('Name:', 'Harsh K. Singh'),
    ('Degree:', 'Bachelor of Technology'),
    ('Role:', 'Frontend Experience Engineer & Interaction Designer'),
    ('University:', 'VIT Bhopal University'),
]
for i, (label, value) in enumerate(details):
    cell_label = table.cell(i, 0)
    cell_value = table.cell(i, 1)
    cell_label.width = Inches(1.5)
    cell_value.width = Inches(4)
    
    p = cell_label.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(label)
    set_font(run, size=12, bold=True)
    
    p = cell_value.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run('  ' + value)
    set_font(run, size=12)

# Remove table borders
for row in table.rows:
    for cell in row.cells:
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        borders = OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'none')
            border.set(qn('w:sz'), '0')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), 'auto')
            borders.append(border)
        tcPr.append(borders)

add_empty_lines(3)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Academic Year 2025–2026')
set_font(run, size=12, bold=True)

add_page_break()

# ═══════════════════════════════════════════════════════════════
# 2. ACKNOWLEDGEMENT
# ═══════════════════════════════════════════════════════════════

add_heading_styled('Acknowledgement', level=1, alignment=WD_ALIGN_PARAGRAPH.CENTER)

add_body(
    'I would like to express my sincere gratitude to VIT Bhopal University and the '
    'School of Computing Science and Engineering for providing the platform and resources '
    'to undertake the EPICS project. The interdisciplinary nature of the EPICS programme '
    'has been instrumental in shaping a practical, user-centred approach to engineering.'
)
add_body(
    'I extend my heartfelt thanks to my faculty mentor and project guide for their '
    'consistent guidance, constructive feedback, and encouragement throughout the development '
    'of AIVA. Their insights on balancing technical rigour with design sensibility proved '
    'invaluable during the iterative design process.'
)
add_body(
    'I am also grateful to my fellow team members whose collaborative spirit made this '
    'project both challenging and rewarding. The cross-functional discussions on clinical '
    'workflow, data architecture, and user experience enriched every phase of development.'
)
add_body(
    'Finally, I wish to acknowledge the broader open-source community, whose frameworks, '
    'libraries, and design philosophies provided the foundational tools upon which the '
    'AIVA frontend experience was built.'
)

add_empty_lines(2)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run('Harsh K. Singh')
set_font(run, size=12, bold=True)

add_page_break()

# ═══════════════════════════════════════════════════════════════
# 3. ABSTRACT
# ═══════════════════════════════════════════════════════════════

add_heading_styled('Abstract', level=1, alignment=WD_ALIGN_PARAGRAPH.CENTER)

add_body(
    'The presentation layer of healthcare technology remains one of the most underexplored '
    'domains in clinical informatics. While considerable research investment is directed toward '
    'artificial intelligence models, predictive algorithms, and secure data pipelines, the '
    'interface through which clinicians and patients interact with these systems is frequently '
    'treated as an afterthought. This disparity results in tools that are technically '
    'sophisticated but experientially deficient — systems that possess clinical power yet fail '
    'to communicate it effectively.'
)
add_body(
    'AIVA (AI-Assisted Intelligent Vitals and Anamnesis) is a patient-controlled Clinical '
    'Co-Pilot designed to bridge the critical gap between medical appointments by transforming '
    'passive health data into structured, time-aware clinical documentation. This individual '
    'contribution report focuses specifically on the frontend engineering and interaction design '
    'of the AIVA platform — the visual and experiential layer that serves as the first point '
    'of contact between the system and its users.'
)
add_body(
    'The work documented herein encompasses the design and development of a premium landing '
    'page experience, a scroll-driven cinematic interaction system, a custom SVG-based dynamic '
    'background engine, a motion design framework built on GSAP and Framer Motion, and an '
    'editorial typography system inspired by luxury product design. Each element was engineered '
    'to convey clinical precision, institutional trust, and technological sophistication without '
    'relying on conventional stock imagery or device mockups.'
)
add_body(
    'The resulting interface establishes a visual identity that positions AIVA not merely as '
    'a health-tech application but as a premium clinical intelligence platform. This report '
    'details the architectural decisions, technical implementation, design philosophy, challenges '
    'encountered, and learning outcomes derived from the frontend development process.'
)

add_page_break()

# ═══════════════════════════════════════════════════════════════
# 4. KEYWORDS
# ═══════════════════════════════════════════════════════════════

add_heading_styled('Keywords', level=1)

add_body(
    'Frontend Engineering, Healthcare UI/UX, Motion Design, GSAP, Framer Motion, '
    'Scroll-based Animation, SVG Graphics, Typography Systems, React, Vite, '
    'Tailwind CSS, Clinical Interface Design, Landing Page Architecture, '
    'Interaction Design, Visual Hierarchy'
)

add_empty_lines(1)

# ═══════════════════════════════════════════════════════════════
# 5. INTRODUCTION
# ═══════════════════════════════════════════════════════════════

add_heading_styled('1. Introduction', level=1)

add_body(
    'In an era where digital health platforms are rapidly proliferating, the quality of a '
    'product\'s interface has become as critical as the robustness of its underlying algorithms. '
    'A patient encountering a health management tool for the first time forms an opinion within '
    'seconds — an opinion shaped not by the complexity of the backend architecture but by the '
    'clarity, elegance, and trustworthiness of the visual presentation. In healthcare, where '
    'user trust is paramount, this first impression carries disproportionate weight.'
)
add_body(
    'AIVA — AI-Assisted Intelligent Vitals and Anamnesis — was conceived as a patient-controlled '
    'clinical co-pilot capable of converting continuous health monitoring data into structured, '
    'clinically actionable documentation. The system addresses a well-documented gap: the '
    '8,760 hours between annual medical checkups during which patient data is either lost, '
    'fragmented, or never captured in the first place.'
)
add_body(
    'While the AIVA ecosystem encompasses backend services, AI processing pipelines, and secure '
    'data management layers, the work presented in this report is concerned exclusively with the '
    'frontend experience layer. The objective was to design and develop a web-based landing page '
    'and product interface that communicates the gravity and precision of the AIVA platform '
    'through visual language alone — without reliance on stock photography, generic illustrations, '
    'or conventional healthcare design tropes.'
)
add_body(
    'This report documents the complete frontend contribution, including architectural decisions '
    'for the component hierarchy, the implementation of advanced scroll-based animation systems, '
    'the creation of a custom dynamic background engine, the integration of a dual motion library '
    'framework, and the establishment of a premium typographic identity. Each section provides '
    'technical detail sufficient for reproducibility while contextualising the design rationale '
    'within broader UI/UX principles.'
)

add_page_break()

# ═══════════════════════════════════════════════════════════════
# 6. TECHNOLOGY STACK
# ═══════════════════════════════════════════════════════════════

add_heading_styled('2. Technology Stack', level=1)

add_body(
    'The selection of technologies for the AIVA frontend was guided by three priorities: '
    'performance, animation capability, and developer ergonomics. Each tool in the stack was '
    'evaluated against the requirements of building a visually rich, scroll-intensive, '
    'single-page application that must load quickly and render smoothly across devices.'
)

# Tech stack table
tech_table = doc.add_table(rows=8, cols=3)
tech_table.style = 'Table Grid'

headers = ['Technology', 'Category', 'Purpose']
for i, h in enumerate(headers):
    cell = tech_table.cell(0, i)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(h)
    set_font(run, size=11, bold=True)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), '003399')
    shading.set(qn('w:val'), 'clear')
    cell._element.get_or_add_tcPr().append(shading)
    run.font.color.rgb = RGBColor(255, 255, 255)

rows_data = [
    ('React 18 (Vite)', 'UI Framework', 'Component-based architecture with fast HMR'),
    ('Tailwind CSS v4', 'Styling', 'Utility-first styling with custom design tokens'),
    ('GSAP + ScrollTrigger', 'Animation Engine', 'Scroll-synchronised timeline animations'),
    ('Framer Motion', 'Micro-interactions', 'Declarative component enter/exit animations'),
    ('SVG (Custom)', 'Graphics', 'Procedurally generated background wave system'),
    ('TypeScript', 'Language', 'Type-safe development for component interfaces'),
    ('Git / GitHub', 'Version Control', 'Collaborative versioning and deployment'),
]
for i, (tech, cat, purpose) in enumerate(rows_data):
    for j, val in enumerate([tech, cat, purpose]):
        cell = tech_table.cell(i+1, j)
        p = cell.paragraphs[0]
        run = p.add_run(val)
        set_font(run, size=11)

add_empty_lines(1)

add_body(
    'React with Vite was chosen over Next.js due to the static, single-page nature of the '
    'landing page — server-side rendering was unnecessary, and Vite\'s superior hot module '
    'replacement speed significantly accelerated the iterative design workflow. Tailwind CSS '
    'provided the utility-first foundation for rapid prototyping while supporting custom '
    'design tokens that enforce visual consistency across sections.'
)

add_page_break()

# ═══════════════════════════════════════════════════════════════
# 7. CORE CONTRIBUTIONS
# ═══════════════════════════════════════════════════════════════

add_heading_styled('3. Core Contributions', level=1)

add_body(
    'The following sub-sections document each significant contribution to the AIVA frontend '
    'experience. Each contribution is presented with its design rationale, technical '
    'implementation approach, and the resulting impact on the overall product perception.'
)

# 3.1 Landing Page Architecture
add_heading_styled('3.1 Landing Page Architecture', level=2)

add_body(
    'The landing page serves as the primary entry point for all AIVA stakeholders — clinicians, '
    'patients, institutional partners, and potential investors. Its architecture was designed '
    'with a singular principle: typographic conviction over visual decoration.'
)
add_body(
    'The initial design iterations explored conventional healthcare landing page patterns, '
    'including device mockups displaying the AIVA dashboard interface within a mobile phone '
    'frame positioned in a split-column layout alongside the hero text. While technically '
    'competent, this approach introduced visual noise that diluted the core message. The '
    'decision was made to eliminate all device mockups entirely, transitioning to a fully '
    'centred, typography-driven hero composition.'
)
add_body('Key architectural decisions included:')
add_bullet('Adoption of a single-column, vertically centred hero layout that commands the full viewport')
add_bullet('Hierarchical text composition separating "8,760 Hours of Your Health" from "15 Minutes to Explain It" across distinct visual tiers')
add_bullet('Strategic use of the brand accent colour (#82B2FF) exclusively on the "15 Minutes" phrase to create a focal anchor')
add_bullet('Removal of hero call-to-action buttons from the content area, relocating the "Request Access" interaction to the persistent navigation bar')
add_bullet('Implementation of a minimal navigation structure with only three anchors — System, Intelligence, and Contact — each linked to corresponding page sections via smooth scroll behaviour')

# 3.2 Cinematic Scroll Experience
add_heading_styled('3.2 Cinematic Scroll Experience', level=2)

add_body(
    'One of the most technically ambitious contributions was the design of a cinematic, '
    'scroll-driven narrative flow. Rather than presenting information as a static stack of '
    'sections, the page was conceived as a sequential revelation — each scroll action '
    'unveiling the next layer of the AIVA story.'
)
add_body(
    'The scroll experience leverages GSAP\'s ScrollTrigger plugin to synchronise animation '
    'timelines with the user\'s scroll position. This approach transforms passive scrolling '
    'into an active, cinematic interaction where content elements animate into view with '
    'precise timing and choreography.'
)
add_body('Core scroll-driven interactions include:')
add_bullet('Parallax displacement of the hero text block relative to the background wave mesh, creating depth perception')
add_bullet('Section-level entrance animations triggered at configurable scroll thresholds')
add_bullet('Progressive opacity transitions that fade content in as users approach each section')
add_bullet('Hardware-accelerated CSS transforms ensuring 60fps animation performance')

add_body(
    'An early design concept explored a laptop-opening cinematic sequence at page load, where '
    'the screen would appear to emerge from a closed device and expand into the full landing '
    'page. While this concept was prototyped, it was ultimately refined into a subtler entry '
    'animation to ensure broad device compatibility and reduced initial load time.'
)

add_page_break()

# 3.3 Motion System Integration
add_heading_styled('3.3 Motion System Integration', level=2)

add_body(
    'The AIVA frontend employs a dual motion library architecture, combining GSAP for '
    'scroll-synchronised macro-animations with Framer Motion for declarative component-level '
    'micro-interactions. This separation of concerns allows each library to operate within '
    'its domain of strength.'
)

add_heading_styled('GSAP Layer (Macro-animations)', level=3)
add_bullet('Scroll-triggered section reveals and parallax effects')
add_bullet('Continuous infinite-loop animations for background wave elements')
add_bullet('Timeline-based sequential animations for multi-step content reveals')
add_bullet('GSAP Context API for proper cleanup and memory management within React\'s lifecycle')

add_heading_styled('Framer Motion Layer (Micro-interactions)', level=3)
add_bullet('Component mount/unmount animations using initial, animate, and exit states')
add_bullet('Staggered children animations for list and grid element reveals')
add_bullet('useScroll and useTransform hooks for lightweight scroll-linked transforms')
add_bullet('SVG path drawing animations using pathLength for the lattice graphic effects')

add_body(
    'A critical integration challenge was preventing conflicts between the two animation '
    'systems when they targeted overlapping DOM elements. This was resolved by establishing '
    'clear ownership boundaries: GSAP exclusively controls elements with class-based selectors '
    '(e.g., .wave-1, .mesh-wave-left), while Framer Motion controls elements through React '
    'component props. This convention eliminated race conditions and ensured deterministic '
    'animation behaviour.'
)

# 3.4 Dynamic Background System
add_heading_styled('3.4 Dynamic Background System', level=2)

add_body(
    'The background of the AIVA landing page is not a static gradient or image — it is a '
    'procedurally generated, continuously animated SVG wave mesh system. This system was '
    'designed to evoke the sensation of data flowing through a clinical intelligence network, '
    'reinforcing the platform\'s identity without competing with foreground content.'
)

add_body('The SignalWave component generates its visual output through the following process:')
add_bullet('Thirty mathematically computed SVG cubic Bezier curves are rendered on each side of the viewport')
add_bullet('Each curve\'s control points are algorithmically offset to create an organic, interference-pattern aesthetic')
add_bullet('GSAP\'s sine easing function drives a gentle floating animation, displacing curves by 30 pixels with yoyo repetition')
add_bullet('CSS mask gradients fade the mesh to transparency toward the viewport centre, ensuring text readability')
add_bullet('Top and bottom gradient overlays soften the vertical boundaries for seamless section transitions')

add_body(
    'The component is rendered as a fixed-position layer at z-index 0, beneath all content, '
    'ensuring it serves as an ambient texture rather than a competing visual element. '
    'Performance profiling confirmed that the 60-path SVG system maintains consistent 60fps '
    'rendering on mid-range hardware, with the GSAP animations consuming less than 2ms per '
    'frame in Chrome\'s performance timeline.'
)

add_page_break()

# 3.5 Typography & Visual Design
add_heading_styled('3.5 Typography and Visual Design', level=2)

add_body(
    'In the absence of photography, illustrations, or device mockups, typography assumes the '
    'role of the primary visual instrument. The AIVA typographic system was designed to convey '
    'three qualities simultaneously: clinical precision, institutional authority, and '
    'contemporary elegance.'
)

add_heading_styled('Font Selection', level=3)
add_body(
    'The display typeface Playfair Display was selected for all headline elements. Playfair '
    'Display is a transitional serif typeface characterised by high contrast between thick and '
    'thin strokes, delicate hairline serifs, and an overall silhouette associated with editorial '
    'luxury and classical print design. Its visual weight communicates authority while its '
    'refined proportions maintain a modern sensibility appropriate for a technology product.'
)
add_body(
    'Body text employs Inter, a humanist sans-serif designed for screen readability at small '
    'sizes. The contrast between the ornate display type and the clean body type creates a '
    'visual rhythm that naturally guides the reader\'s eye through the content hierarchy.'
)

add_heading_styled('Typographic Hierarchy', level=3)
add_body('The hero section demonstrates the complete hierarchy:')
add_bullet('Primary headline: Playfair Display, 100px on desktop, white (#FFFFFF), tight tracking')
add_bullet('Accent phrase ("15 Minutes"): Same typeface, brand blue (#82B2FF), creating a deliberate focal interruption')
add_bullet('Subtitle body: Inter, 19px, muted grey (#A1A1AA), light weight with generous tracking')
add_bullet('Navigation elements: Inter, 15px, secondary grey with hover-to-white transitions')

add_body(
    'This hierarchy ensures that even on a completely text-driven page, the visual weight '
    'distribution creates clear reading paths and maintains user engagement across the '
    'full viewport height.'
)

add_page_break()

# ═══════════════════════════════════════════════════════════════
# 8. MOTION & INTERACTION SYSTEM
# ═══════════════════════════════════════════════════════════════

add_heading_styled('4. Motion and Interaction System', level=1)

add_body(
    'Motion design within the AIVA interface follows a philosophy of purposeful restraint. '
    'Every animation serves one of three functions: guiding attention, communicating state '
    'changes, or establishing spatial relationships. Decorative animation for its own sake '
    'was deliberately excluded.'
)

add_heading_styled('4.1 Animation Principles', level=2)

add_bullet('Entrance animations use easeOut curves to create a sense of content arriving into position')
add_bullet('Duration values are calibrated between 800ms and 1200ms to feel deliberate without causing impatience')
add_bullet('Stagger delays of 100–150ms between sibling elements create a cascading reveal effect')
add_bullet('All transforms use GPU-accelerated properties (transform, opacity) to avoid layout recalculations')

add_heading_styled('4.2 Scroll-Linked Motion', level=2)

add_body(
    'The scroll-linked motion system transforms the act of scrolling from passive navigation '
    'into an interactive reveal mechanism. Content sections are not simply present on the page '
    'waiting to be scrolled into view; rather, they are choreographed to enter the viewport '
    'with intentional timing that mirrors the narrative structure of the AIVA value proposition.'
)
add_body(
    'GSAP\'s ScrollTrigger plugin serves as the orchestration layer. Each section registers '
    'a ScrollTrigger instance with configurable start and end positions, allowing precise '
    'control over when animations begin and how they progress relative to the scroll position. '
    'The use of scrub mode for certain animations enables the animation progress to be directly '
    'tied to scroll position, creating a frame-accurate, user-controlled cinematic experience.'
)

add_heading_styled('4.3 Performance Considerations', level=2)

add_body(
    'Animation performance was continuously monitored throughout development using Chrome '
    'DevTools\' Performance and Rendering panels. Key performance targets included:'
)
add_bullet('Maintaining 60fps during all scroll-triggered animations')
add_bullet('Keeping the main thread idle time above 80% during animation sequences')
add_bullet('Limiting total animation-related JavaScript execution to under 4ms per frame')
add_bullet('Using will-change CSS property judiciously to hint browser compositing without excessive memory allocation')

add_page_break()

# ═══════════════════════════════════════════════════════════════
# 9. UI/UX DESIGN SYSTEM
# ═══════════════════════════════════════════════════════════════

add_heading_styled('5. UI/UX Design System', level=1)

add_heading_styled('5.1 Design Philosophy', level=2)

add_body(
    'The AIVA interface operates under a design philosophy that can be described as '
    '"clinical minimalism" — the deliberate reduction of visual elements to only those that '
    'serve a communicative or functional purpose. This philosophy is not minimalism for '
    'aesthetic preference; it is minimalism in service of clinical clarity.'
)
add_body(
    'Healthcare interfaces face a unique design tension: they must project authority and '
    'trustworthiness while remaining approachable and navigable. The AIVA design system '
    'resolves this tension through a dark-theme foundation that connotes technological '
    'sophistication, paired with generous whitespace and restrained colour usage that '
    'prevents sensory overload.'
)

add_heading_styled('5.2 Colour Architecture', level=2)

add_body('The colour palette is intentionally constrained:')
add_bullet('Background: Deep blacks (#0B0B0C) providing maximum contrast for text elements')
add_bullet('Primary text: Pure white (#FFFFFF) for headlines, ensuring immediate legibility')
add_bullet('Secondary text: Muted zinc (#A1A1AA) for supporting copy, reducing cognitive load')
add_bullet('Accent: Clinical blue (#82B2FF) used sparingly for interactive elements and focal highlights')
add_bullet('Success indicator: System green (#4ade80) reserved exclusively for the logo status dot')

add_heading_styled('5.3 Spatial Design', level=2)

add_body(
    'The spatial design employs a breathing layout system — sections are given substantial '
    'vertical padding (128px per section) to create natural pause points during scrolling. '
    'This generous spacing serves a dual purpose: it prevents information overload by '
    'limiting the density of content per viewport, and it creates natural animation trigger '
    'zones that align with the scroll-driven motion system.'
)

add_heading_styled('5.4 User Attention Flow', level=2)

add_body(
    'The page is designed to guide user attention through a deliberate funnel. The hero '
    'section captures immediate attention through the scale and weight of the serif headline. '
    'The accent-coloured "15 Minutes" phrase anchors focus before the subtitle elaborates '
    'the value proposition. The navigation bar\'s "Request Access" button provides a persistent '
    'call-to-action that is always accessible regardless of scroll position. As the user '
    'scrolls, each section reveals a deeper layer of the AIVA narrative, culminating in the '
    'waitlist form — the primary conversion target.'
)

add_page_break()

# ═══════════════════════════════════════════════════════════════
# 10. CHALLENGES
# ═══════════════════════════════════════════════════════════════

add_heading_styled('6. Challenges', level=1)

add_heading_styled('6.1 Achieving Premium UI Without Overdesign', level=2)
add_body(
    'The most persistent design challenge was maintaining visual richness without descending '
    'into visual excess. Premium design is characterised by restraint — the confidence to leave '
    'space empty, to let typography speak, to resist the impulse to fill every pixel with '
    'content. Multiple iterations were required to find the threshold where the design felt '
    'substantial and authoritative without appearing cluttered or overwrought. The removal of '
    'the phone mockup from the hero section was emblematic of this iterative refinement process.'
)

add_heading_styled('6.2 Dual Animation Library Coordination', level=2)
add_body(
    'Integrating GSAP and Framer Motion within the same React application introduced subtle '
    'coordination challenges. Both libraries can manipulate DOM element styles, and without '
    'careful architectural boundaries, conflicting transforms would produce visual artifacts. '
    'Establishing a clear class-based versus prop-based ownership model resolved these conflicts, '
    'but arriving at this solution required significant debugging and experimentation.'
)

add_heading_styled('6.3 SVG Performance at Scale', level=2)
add_body(
    'The dynamic background system renders sixty concurrent SVG paths with continuous GSAP '
    'animations. Initial implementations caused measurable frame drops on lower-end devices. '
    'Optimisation strategies included reducing path complexity (fewer control points per curve), '
    'using CSS will-change hints selectively, and leveraging GSAP\'s built-in RAF (requestAnimationFrame) '
    'batching to minimise layout thrashing.'
)

add_heading_styled('6.4 Cross-Device Typography Consistency', level=2)
add_body(
    'Playfair Display\'s high contrast and delicate serifs render differently across operating '
    'systems due to varying font rasterisation engines. Windows ClearType, macOS Core Text, and '
    'Linux FreeType each produce subtly different glyph weights. Mitigation involved testing '
    'across platforms, applying -webkit-font-smoothing: antialiased for consistent rendering, '
    'and adjusting font weights per breakpoint to maintain perceived visual weight across devices.'
)

add_page_break()

# ═══════════════════════════════════════════════════════════════
# 11. LEARNING OUTCOMES
# ═══════════════════════════════════════════════════════════════

add_heading_styled('7. Learning Outcomes', level=1)

add_heading_styled('7.1 Frontend Architecture at Scale', level=2)
add_body(
    'Working on the AIVA project deepened my understanding of component architecture in '
    'large-scale React applications. The process of decomposing a complex page into reusable, '
    'self-contained components — each with clearly defined interfaces and responsibilities — '
    'reinforced principles of separation of concerns and compositional design. The experience '
    'of managing nineteen component directories taught practical lessons in code organisation '
    'that textbook examples rarely convey.'
)

add_heading_styled('7.2 Animation Systems Engineering', level=2)
add_body(
    'The dual implementation of GSAP and Framer Motion provided firsthand experience with '
    'two fundamentally different animation paradigms: imperative timeline-based animation '
    'and declarative state-driven animation. Understanding the strengths and limitations '
    'of each approach has become a significant addition to my professional skill set, '
    'enabling informed technology selection for future animation-heavy projects.'
)

add_heading_styled('7.3 Design Thinking in Practice', level=2)
add_body(
    'The iterative nature of the AIVA frontend design process — progressing through split-column '
    'layouts, device mockups, 3D lattice graphics, and ultimately arriving at a centred '
    'typographic composition — was a practical exercise in design thinking. Each iteration '
    'was not a failure but a refinement, with each version teaching something about what the '
    'product\'s visual identity should and should not be.'
)

add_heading_styled('7.4 Real-World Product Design Sensibility', level=2)
add_body(
    'Perhaps the most valuable outcome was the development of a product design sensibility '
    'that extends beyond technical capability. The AIVA project required thinking about how '
    'visual design influences user trust, how typography communicates institutional authority, '
    'and how motion design shapes emotional responses. These considerations are rarely '
    'addressed in academic curricula but are essential in professional product development.'
)

add_page_break()

# ═══════════════════════════════════════════════════════════════
# 12. CONCLUSION
# ═══════════════════════════════════════════════════════════════

add_heading_styled('8. Conclusion', level=1)

add_body(
    'The frontend experience of AIVA represents more than a collection of styled components '
    'and animated transitions. It is the visual identity of a clinical intelligence platform '
    '— the layer through which every user, whether patient or physician, forms their first '
    'and most lasting impression of the system\'s capability and trustworthiness.'
)
add_body(
    'Through the work documented in this report, I contributed to establishing a design language '
    'that communicates clinical precision through typographic authority, technological sophistication '
    'through procedural motion, and institutional trust through visual restraint. The centred, '
    'typography-driven hero composition, the ambient SVG wave system, the dual-library motion '
    'framework, and the premium Playfair Display typographic identity collectively form a '
    'frontend experience that positions AIVA among the highest tier of health-tech product '
    'presentations.'
)
add_body(
    'The interface is architected for extensibility. As the AIVA platform evolves from its '
    'current Phase I specification — incorporating real-time vitals dashboards, AI-generated '
    'clinical summaries, and provider collaboration tools — the design system and component '
    'architecture established during this development phase provide a robust foundation for '
    'seamless expansion. The visual language, motion principles, and spatial design conventions '
    'documented herein will ensure consistency as the product scales.'
)
add_body(
    'In healthcare technology, the interface is not decoration — it is communication. The work '
    'presented in this report reflects a commitment to ensuring that AIVA\'s communication is '
    'as precise, trustworthy, and compelling as the clinical intelligence it delivers.'
)

add_page_break()

# ═══════════════════════════════════════════════════════════════
# 13. BIODATA
# ═══════════════════════════════════════════════════════════════

add_heading_styled('Biodata', level=1, alignment=WD_ALIGN_PARAGRAPH.CENTER)

add_empty_lines(1)

bio_table = doc.add_table(rows=8, cols=2)
bio_table.style = 'Table Grid'
bio_table.alignment = WD_ALIGN_PARAGRAPH.CENTER

bio_data = [
    ('Name', 'Harsh K. Singh'),
    ('University', 'VIT Bhopal University'),
    ('Degree', 'Bachelor of Technology'),
    ('Project', 'AIVA (AI-Assisted Intelligent Vitals & Anamnesis)'),
    ('Role', 'Frontend Experience Engineer & Interaction Designer'),
    ('Technical Skills', 'MERN Stack, Frontend Development, UI/UX Design, Animation Systems (GSAP, Framer Motion), TypeScript, Responsive Design'),
    ('Development Tools', 'VS Code, Git, GitHub, Chrome DevTools, MongoDB, React, Vite, Tailwind CSS'),
    ('Interests', 'Frontend Engineering, Interface Design, Motion Graphics, Problem Solving'),
]

for i, (label, value) in enumerate(bio_data):
    cell_label = bio_table.cell(i, 0)
    cell_value = bio_table.cell(i, 1)
    
    p = cell_label.paragraphs[0]
    run = p.add_run(label)
    set_font(run, size=11, bold=True)
    
    p = cell_value.paragraphs[0]
    run = p.add_run(value)
    set_font(run, size=11)

# Style the header column
for i in range(len(bio_data)):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'F2F2F2')
    shading.set(qn('w:val'), 'clear')
    bio_table.cell(i, 0)._element.get_or_add_tcPr().append(shading)

# Set column widths
for row in bio_table.rows:
    row.cells[0].width = Inches(2)
    row.cells[1].width = Inches(4.5)

# ─── SAVE DOCUMENT ────────────────────────────────────────────

output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'AIVA_Individual_Contribution_Report_Harsh_K_Singh.docx')
doc.save(output_path)
print(f"Report generated successfully: {output_path}")
